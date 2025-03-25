from docx import Document
from docx.text.paragraph import Paragraph
from translate import translate, validate, SqlWrapper
from utils import load_config
from loguru import logger
import string
import re
import argparse
import os

# Load the configuration
config = load_config()


def is_title(paragraph):
    # List of words that are not usually capitalized in a title
    exceptions = {
        "a",
        "an",
        "and",
        "as",
        "at",
        "but",
        "by",
        "for",
        "in",
        "nor",
        "of",
        "on",
        "or",
        "the",
        "up",
    }

    # Remove punctuation using str.translate and string.punctuation
    translator = str.maketrans("", "", string.punctuation)
    cleaned_paragraph = paragraph.translate(translator)

    # Check if the cleaned paragraph is all uppercase
    if cleaned_paragraph.replace(" ", "").isupper():
        return True

    # Split the cleaned paragraph into words
    words = cleaned_paragraph.split()

    # Check if the cleaned paragraph is likely a title based on the capitalization
    for i, word in enumerate(words):
        word = re.sub(r"[^A-Za-z0-9]+", "", word)
        if i == 0 or i == len(words) - 1:
            if not word.istitle() and not word.isupper():
                return False
        elif word.lower() not in exceptions:
            if not word.istitle() and not word.isupper():
                return False
        else:
            if word.isupper() or (word.istitle() and word.lower() in exceptions):
                return False

    return True


def is_bold(paragraph):
    for run in paragraph.runs:
        if not run.bold:
            return False
    return True


def get_style(paragraph):
    longest_run = None
    max_length = 0
    for run in paragraph.runs:
        if len(run.text) > max_length:
            longest_run = run
            max_length = len(run.text)
    if longest_run:
        return longest_run.style, longest_run.font.size
    else:
        return None


def add_text_to_paragraph(paragraph, new_text, translation_only=False):
    """
    Add text to a given Paragraph object, using the style and font size of the run
    with the highest character count.

    Parameters:
    - paragraph: A docx.text.paragraph.Paragraph object.
    - new_text: The text to add to the paragraph.
    - translation_only: If True, replace the original text instead of appending
    """
    if new_text.isdigit():
        return

    if not isinstance(paragraph, Paragraph):
        raise TypeError(
            "The provided paragraph must be a docx.text.paragraph.Paragraph object."
        )

    result = get_style(paragraph)

    if translation_only:
        # Clear existing runs
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    if result:
        style, font_size = result
        new_run = paragraph.add_run(new_text)
        new_run.style = style
        if font_size:
            new_run.font.size = font_size
        if is_bold(paragraph):
            new_run.bold = True
    else:
        paragraph.add_run(new_text)


def is_page_number(paragraph):
    # If a paragraph begins/ends by numbers and < 10 words, without period, it's likely a page number
    first_word = paragraph.text.strip().split()[0]
    last_word = paragraph.text.strip().split()[-1]
    if (
        (first_word.isdigit() or last_word.isdigit())
        and len(paragraph.text.split()) < 10
        and "." not in paragraph.text
    ):
        return True
    return False


def process_paragraphs(doc):
    """
    Process paragraphs and return information about which paragraphs to translate
    and how they're connected.
    """
    last_char = ""
    prev_style = None
    prev_paragraph = None
    paragraph_maps = {}
    final_paragraphs = set()

    for i, paragraph in enumerate(doc.paragraphs):
        style = get_style(paragraph)
        if paragraph.text.strip() == "":
            continue
        elif (
            is_title(paragraph.text.strip())
            or is_page_number(paragraph)
            or is_bold(paragraph)
        ):
            final_paragraphs.add(i)
            continue
        elif (
            (last_char.isalnum() or last_char == ",")
            and (paragraph.text.strip()[0].isalnum())
        ) and style == prev_style:
            last_char = paragraph.text.strip()[-1]
            paragraph_maps[i] = prev_paragraph
            prev_paragraph = i
        else:
            last_char = paragraph.text.strip()[-1]
            prev_style = get_style(paragraph)
            if prev_paragraph:
                final_paragraphs.add(prev_paragraph)
            prev_paragraph = i
            paragraph_maps[i] = None

    if prev_paragraph:
        final_paragraphs.add(prev_paragraph)

    return final_paragraphs, paragraph_maps


def translate_doc(docx_filename, output_filename, args, translation_only=False):
    """
    Translate a document and save the output.

    Parameters:
    - docx_filename: Input document filename
    - output_filename: Output document filename
    - args: Command line arguments
    - translation_only: If True, only include translations in output
    """
    # Load the document
    doc = Document(docx_filename)

    final_paragraphs, paragraph_maps = process_paragraphs(doc)

    with SqlWrapper(f'output/{config["CN_TITLE"]}/buffer.db') as cache:
        for i, p in enumerate(doc.paragraphs):
            if i not in final_paragraphs:
                continue
            text_to_translate = p.text.strip()
            while i in paragraph_maps and paragraph_maps[i]:
                text_to_translate = (
                    doc.paragraphs[paragraph_maps[i]].text.strip()
                    + " "
                    + text_to_translate
                )
                i = paragraph_maps[i]
            if text_to_translate in cache and validate(
                text_to_translate, cache[text_to_translate]
            ):
                translated_text = cache[text_to_translate]
            elif text_to_translate.isdigit():
                continue
            else:
                translated_text = translate(text_to_translate, dryrun=args.dryrun)
                if not args.dryrun:
                    cache[text_to_translate] = translated_text

            if translation_only:
                add_text_to_paragraph(p, translated_text, translation_only=True)
            else:
                add_text_to_paragraph(p, "\n" + translated_text)

    doc.save(output_filename)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--dryrun", action="store_true")
    args = parser.parse_args()

    if args.dryrun:
        logger.warning("Dry run mode enabled. No translation will be performed.")

    output_dir = f'output/{config["CN_TITLE"]}'
    logger.add(f"{output_dir}/info.log", colorize=True, level="DEBUG")

    # Generate both dual-language and translation-only outputs
    translate_doc(
        f"{output_dir}/input.docx", f'{output_dir}/{config["CN_TITLE"]}_dual.docx', args
    )
    translate_doc(
        f"{output_dir}/input.docx",
        f'{output_dir}/{config["CN_TITLE"]}_translated.docx',
        args,
        translation_only=True,
    )
